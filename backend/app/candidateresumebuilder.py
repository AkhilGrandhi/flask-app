# resume_blueprint.py
from flask import Blueprint, request, send_file, jsonify
from openai import OpenAI
from io import BytesIO
import os
import re
import traceback
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

# ------- Word (python-docx) -------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ------- PDF (LibreOffice conversion) -------
import tempfile
import subprocess
import platform

# ------- Database imports -------
from app.models import db, CandidateJob

bp = Blueprint("resume", __name__)

# ---- Config: Load API key from environment variable ----
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# ---- Section detection ----
SECTION_TITLES = {
    "professional summary",
    "summary",
    "technical skills",
    "skills",
    "professional experience",
    "experience",
    "work experience",
    "work history",
    "education",
    "certifications",
    "projects",
    "additional qualifications",
    "additional information",
    "references",
}

def clean_markdown(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = text.replace("`", "")
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = text.replace("**", "").replace("*", "").replace("_", "")
    text = re.sub(r"^\s*[•\-–]\s*", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def is_contact_line(line: str) -> bool:
    if not line:
        return False
    l = line.lower()
    return (
        "email" in l
        or "@" in l
        or "phone" in l
        or re.search(r"\b\d{10}\b", l) is not None
        or re.search(r"\+\d", l) is not None
    )

def is_section_title(line: str) -> bool:
    if not line:
        return False
    raw = line.strip().rstrip(":")
    return raw.lower() in SECTION_TITLES

def add_horizontal_rule(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---- Word building helpers ----
def add_candidate_name(doc, lines, idx):
    if idx < len(lines):
        name_para = doc.add_paragraph(lines[idx])
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_para.runs[0]
        run.bold = True
        run.font.size = Pt(20)
        idx += 1
    return idx

def add_contact_info(doc, lines, idx):
    contact_email, contact_phone, contact_location = "", "", ""
    while idx < len(lines) and is_contact_line(lines[idx]):
        line = lines[idx]
        email_match = re.search(r"[\w\.-]+@[\w\.-]+", line)
        if email_match:
            contact_email = email_match.group(0)
        phone_match = re.search(r"(\+?\d[\d\s\-]{8,}\d)", line)
        if phone_match:
            contact_phone = phone_match.group(0).strip()
        loc_match = re.search(r"Location\s*[:\-]?\s*(.*)", line, re.IGNORECASE)
        if loc_match:
            contact_location = loc_match.group(1).strip()
        idx += 1

    if contact_email or contact_phone or contact_location:
        pieces = []
        if contact_email:
            pieces.append(f"Email: {contact_email}")
        if contact_phone:
            pieces.append(f"Mobile: {contact_phone}")
        if contact_location:
            pieces.append(f"Location: {contact_location}")
        contact_line = "  |  ".join(pieces)
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para.runs[0].font.size = Pt(11)

    return idx

def add_section_title(doc, title, idx):
    p = doc.add_paragraph(title.upper().rstrip(":"))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)
    add_horizontal_rule(p)
    return idx + 1

def add_skills_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    category = None
    skills = []
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue

        # Case 1: Inline list under an existing category
        if category and not line.startswith("-") and "," in line:
            skills = [s.strip() for s in line.split(",") if s.strip()]
            p = doc.add_paragraph()
            r1 = p.add_run(category + ": ")
            r1.bold = True
            p.add_run(", ".join(skills))
            category, skills = None, []

        # Case 2: New category line
        elif not line.startswith("-"):
            if category and skills:
                p = doc.add_paragraph()
                r1 = p.add_run(category + ": ")
                r1.bold = True
                p.add_run(", ".join(skills))
            category = line
            skills = []

        # Case 3: Bulleted skill
        else:
            skills.append(line.lstrip("- ").strip())
        idx += 1

    # Flush last category
    if category and skills:
        p = doc.add_paragraph()
        r1 = p.add_run(category + ": ")
        r1.bold = True
        p.add_run(", ".join(skills))
    return idx

def add_experience_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    company_seen = False
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx]

        # Company – Location OR Role – Dates
        if " – " in line and ":" not in line:
            if " to " in line:  # role line
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(10)
            else:  # company line
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(11)
                if company_seen:
                    p.paragraph_format.space_before = Pt(10)
                company_seen = True

        elif " – " in line and ":" in line:  # job + bullets in same line
            job_title, rest = line.split(":", 1)
            p = doc.add_paragraph(job_title.strip())
            p.runs[0].bold = True
            parts = re.split(r'\.\s+|,\s+', rest)
            for part in parts:
                if part.strip():
                    bullet_para = doc.add_paragraph(part.strip(), style="List Bullet")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)

        elif line.startswith("Technologies Used"):
            heading, _, techs = line.partition(":")
            p = doc.add_paragraph()
            r1 = p.add_run(heading.strip() + ": ")
            r1.bold = True
            p.add_run(techs.strip())
            p.paragraph_format.space_after = Pt(10)

        elif line.startswith("- "):  # standard bullets
            bullet_para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            bullet_para.paragraph_format.left_indent = Inches(0.25)
        else:
            doc.add_paragraph(line)
        idx += 1
    return idx

def add_certifications_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].lstrip("- ").strip()
        if line:
            doc.add_paragraph(line, style="List Bullet")
        idx += 1
    return idx

def add_education_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    while idx < len(lines) and not is_section_title(lines[idx]):
        doc.add_paragraph(lines[idx])
        idx += 1
    return idx

def add_summary_section(doc, lines, idx):
    idx = add_section_title(doc, lines[idx], idx)
    while idx < len(lines) and not is_section_title(lines[idx]):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        text = line[2:].strip() if line.startswith("- ") else line
        bullet_para = doc.add_paragraph(text, style="List Bullet")
        bullet_para.paragraph_format.left_indent = Inches(0.25)
        idx += 1
    return idx

def extract_total_experience(candidate_info: str) -> str:
    candidate_info = candidate_info.replace("–", "-").replace("—", "-")
    duration_lines = re.findall(r"Duration:\s*(.+)", candidate_info, re.IGNORECASE)
    total_months = 0
    today = datetime.today()

    for line in duration_lines:
        parts = [p.strip() for p in line.split("-")]
        if len(parts) != 2:
            continue
        start_str, end_str = parts
        # parse start
        try:
            start_date = datetime.strptime(start_str, "%b %Y")
        except ValueError:
            start_date = datetime.strptime(start_str, "%B %Y")
        # parse end
        if "present" in end_str.lower():
            end_date = today
        else:
            try:
                end_date = datetime.strptime(end_str, "%b %Y")
            except ValueError:
                end_date = datetime.strptime(end_str, "%B %Y")
        months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
        total_months += months

    years, m = divmod(total_months, 12)
    return f"Total Experience: {years} years {m} months"

# ---- Word generator ----
def create_resume_word(content: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    para_format = style.paragraph_format
    para_format.space_after = Pt(0)
    para_format.space_before = Pt(0)
    para_format.line_spacing = 1
    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    lines = [ln.strip("• ").strip() for ln in content.splitlines() if ln and str(ln).strip()]
    idx = 0

    idx = add_candidate_name(doc, lines, idx)
    idx = add_contact_info(doc, lines, idx)

    while idx < len(lines):
        if is_section_title(lines[idx]):
            section_key = lines[idx].strip().rstrip(":").lower()
            if section_key in ("professional summary", "summary"):
                idx = add_summary_section(doc, lines, idx)
            elif section_key in ("skills", "technical skills"):
                idx = add_skills_section(doc, lines, idx)
            elif section_key in ("work experience", "professional experience"):
                idx = add_experience_section(doc, lines, idx)
            elif section_key == "certifications":
                idx = add_certifications_section(doc, lines, idx)
            elif section_key == "education":
                idx = add_education_section(doc, lines, idx)
            else:
                idx = add_section_title(doc, lines[idx], idx)
        else:
            idx += 1
    return doc

def create_resume_pdf(resume_text: str) -> BytesIO:
    # 1) Create a DOCX via the same builder
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = create_resume_word(resume_text)
    doc.save(tmp_docx.name)

    # 2) Choose LibreOffice executable
    system = platform.system()
    soffice_path = (
        r"C:\Program Files\LibreOffice\program\soffice.exe" if system == "Windows" else "libreoffice"
    )

    # 3) Convert DOCX -> PDF
    try:
        subprocess.run(
            [soffice_path, "--headless", "--convert-to", "pdf", tmp_docx.name, "--outdir", os.path.dirname(tmp_docx.name)],
            check=True
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {e}")
    except FileNotFoundError:
        raise RuntimeError(f"LibreOffice not found at {soffice_path}. Install it or update the path.")

    tmp_pdf_path = os.path.splitext(tmp_docx.name)[0] + ".pdf"
    with open(tmp_pdf_path, "rb") as f:
        pdf_bytes = BytesIO(f.read())

    # 4) Cleanup temps
    try:
        os.remove(tmp_docx.name)
        os.remove(tmp_pdf_path)
    except OSError:
        pass

    pdf_bytes.seek(0)
    return pdf_bytes

# ----------------------------
# Helpers for merging sections
# ----------------------------
def _ensure_title(text: str) -> str:
    return re.sub(r"\b(work experience)\b", "WORK EXPERIENCE", text, flags=re.IGNORECASE)

# ---- API endpoint (Blueprint) ----
@bp.post("/generate")
def generate_resume():
    # Inputs
    try:
        data = request.get_json(force=True, silent=False)
    except Exception:
        return jsonify({"message": "Invalid JSON"}), 400

    job_desc = (data or {}).get("job_desc", "").strip()
    candidate_info = (data or {}).get("candidate_info", "").strip()
    file_type = (data or {}).get("file_type", "word").strip().lower()
    candidate_id = (data or {}).get("candidate_id")
    job_row_id = (data or {}).get("job_row_id")

    if not job_desc or not candidate_info:
        return jsonify({"message": "Missing required fields"}), 400

    work_exp_str = extract_total_experience(candidate_info)

    # Validate OpenAI API key
    if not OPENAI_API_KEY:
        return jsonify({"message": "OpenAI API key not configured. Please set OPENAI_API_KEY environment variable."}), 500
    
    # OpenAI client with timeout
    try:
        client = OpenAI(api_key=OPENAI_API_KEY, timeout=120.0)
    except Exception as e:
        return jsonify({"message": f"OpenAI client init error: {e}"}), 500

    # --- Prompts from the new script (main sections + experience), run in parallel ---
    def generate_main_sections():
        prompt = f"""
            You are a professional resume writer. Using the Job Description and Candidate Information provided below, generate a clean, ATS-optimized resume that strictly follows the section order and formatting rules listed here:

            ⚠️ IMPORTANT: Output must contain the **resume only** — do not include explanations, disclaimers, notes, or extra text outside of the resume.

            SECTION ORDER:

            1. **PROFESSIONAL SUMMARY** – Generate **6 to 8 bullet points**.  
                - The **first bullet point** must always mention the candidate's **total years of professional experience**. If this information is present in the JOB DESCRIPTION, use the role mentioned there when framing the experience.  
                    WORK EXPERIENCE: {work_exp_str}  
                - Represent the total experience as **"X+ years of experience"** (e.g., *5+ years*, *6+ years*).  
                - Each bullet point must be **at least 2 lines long**, providing rich, detailed information. Avoid short or generic bullets.  
                - The **remaining bullet points** (6–8 total) should comprehensively highlight the candidate’s **key skills, achievements, career milestones, and qualifications** that align closely with the given Job Description.  
                - Each bullet must **start with "- "** (a hyphen followed by a space).  
  
            2. **SKILLS** – Based on the Job Description and Candidate Information:

                1. Identify the **most relevant role/position** (e.g., .NET Developer, Java Backend Engineer, Salesforce Developer, Data Engineer, DevOps Engineer).
                2. Create a **resume-ready Skills section** with **10–12 subsections**, tailored to that role and the JD.

                ⚠️ RULES:
                - Subsections must be **category-based** and recruiter-friendly (e.g., Programming Languages, Frameworks & Libraries, Databases, Cloud Platforms, DevOps & CI/CD, Testing & QA, Security & Compliance, Monitoring & Observability, Collaboration Tools).
                - Use concise, ATS-optimized, professional wording for subsection titles.
                - Fill each subsection with **8–20 related technologies/tools**, directly matching the JD and candidate info.
                - Where possible, **expand categories with specific services or tools** (e.g., list AWS services like EC2, S3, Glue, Lambda, CloudWatch — not just "AWS").
                - Always mirror exact JD keywords (e.g., if JD says "GCP, Spark, BigQuery, Kafka" → those must appear under correct categories).
                - Include versions where impactful (e.g., Java 11/17, .NET 6/7, Spring Boot 3.x, Hadoop 3.x).
                - Do not invent irrelevant categories or mix unrelated technologies into the wrong subsection.
                - Always include these **mandatory baseline categories**, even if not explicitly in the JD:
                    - Programming Languages  
                    - Operating Systems  
                    - Cloud Platforms
                    - DevOps & CI/CD Tools  
                    - Development Tools                   

                Example subsections (adjust dynamically per JD):  
                - Programming Languages  
                - Frameworks & Libraries  
                - Databases & Data Warehousing  
                - Big Data & Streaming  
                - Cloud Platforms  
                - DevOps & CI/CD Tools  
                - Testing & QA  
                - Security & Compliance  
                - Monitoring & Observability  
                - Collaboration Tools  
                - Documentation Tools  
                - Operating Systems  

                ⚠️ Ensure each subsection is **fully loaded with at least 8 skills** and contains **16–20 skills where possible**.
                ⚠️ All technologies listed here must also appear in the **Technologies Used** lines under the WORK EXPERIENCE section.



            3. **CERTIFICATIONS**

            4. **EDUCATION** 
                Format the education section clearly and consistently using the structure shown below. 

                Example Format:
                    MS in Computer Science
                    University of XYZ, USA | GPA: 3.8/4.0
                    B.Tech in Computer Science Engineering
                    JNTU Hyderabad | Percentage: 85%

                Make sure the formatting follows this structure exactly:
                [Degree] in [Field of Study]
                [University Name] | [GPA or Percentage]

                Do not include additional details like thesis titles, coursework, or graduation years unless specifically asked.
            
            ⚠️ IMPORTANT: Do NOT generate the WORK EXPERIENCE section. It will be added separately.

            FORMATTING RULES:
            - Display the candidate's **Name** at the top.
            - Center **Email**, **Phone Number**, and **Candidate Location** on the same line directly below the name, using the format:  
            Email: | Mobile: | Location:
            - Use 0.5-inch page margins.
            - Add a tab space before each bullet point.
            - Do not use markdown or bullet characters like "-", "*", or "•".
            - The **SKILLS** section must always follow the defined categories above—never as a plain list.
            - Always ensure the final resume spans at least 2 full pages of Word or PDF output.

            JOB DESCRIPTION:
            {job_desc}

            CANDIDATE INFORMATION:
            {candidate_info}
            """
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You write polished, ATS-friendly resumes."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )
        return resp.choices[0].message.content or ""

    def generate_work_experience():
        exp_prompt = f"""
            Generate ONLY the WORK EXPERIENCE section for this resume.

            3. **WORK EXPERIENCE** – Merge **Work History** and **Work Experience** into a unified section. For each job role:
                - ⚠️ IMPORTANT: Use WORK EXPERIENCE from the CANDIDATE INFORMATION only
                - Include the Job Title, Company Name (bold), Job Location, and timeline using the format:
                    [Company Name] – [Job Location]  
                    [Job Title] – [Start Month Year] to [End Month Year]

            - Add 10 to 15 high-impact bullet points per role. Each bullet point must:
            - Each bullet point must be exactly 2 lines long, with rich and specific details — including technologies used, metrics, project outcomes, team collaboration, challenges faced, and business impact.
            - "When generating points for each company, first identify the industry it operates in, and then tailor the points to be relevant to that specific industry projects.
            - Start with a strong action verb (e.g., Spearheaded, Engineered, Optimized, Automated, Delivered).
            - Focus on achievements, measurable outcomes, and business value rather than just responsibilities.
            - Include quantifiable results wherever possible (e.g., improved ETL performance by 35%, reduced deployment time by 40%, cut costs by 20% annually).
            - Highlight leadership, innovation, automation, and cross-functional collaboration.
            - Showcase modern practices (e.g., Cloud Migration, DevOps, CI/CD automation, Data Engineering, AI/ML, Security, Scalability).
            - Be specific, technical, and results-driven — not generic.
            

            - ⚠️ Validate technology usage against the job timeline:
            - ONLY include technologies, tools, frameworks, or platforms that were **publicly available and in practical use** during the given employment period.
            - Example: Do NOT include Generative AI, Azure OpenAI, MS Fabric, or other technologies launched post-2021 in roles dated 2020 or earlier.
            - Ensure all technologies and practices mentioned are **realistically applicable** based on release year and industry adoption timeline.

            - Total bullet points should follow this logic:
            - For 1 company: 15 to 20 bullet points.
            - For 2 companies: 15 to 20 bullet points each (total: 30-40 points).
            - For 3 companies: 10 to 15 bullet points each (total: 30-45 points).
            - For 4 companies: 10 to 15 bullet points each (total: 40-60 points).
            - For 5 companies: 10 to 15 bullet points each (total: 60-70 points).
            - For 6 companies: 10 to 15 bullet points each (total: 70-80 points).
            - For 7 companies: 10 to 15 bullet points each (total: 70-80 points).

            - No filler or repetition: Each bullet point must offer unique, concrete contributions or achievements.

            - Write in professional resume tone, use strong action verbs, and focus on clarity, impact, and relevance to technical or engineering roles.

            - End each job section with the line:  
            Technologies Used: tech1, tech2, ..., tech15  
                ⚠️ Ensure each role includes 10 to 15 technologies mapped directly from the SKILLS section.  
                ⚠️ Across all roles, the union of technologies must comprehensively cover the entire SKILLS section.

            JOB DESCRIPTION:
            {job_desc}

            CANDIDATE INFORMATION:
            {candidate_info}
            """
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You write only the Work Experience section for ATS resumes."},
                {"role": "user", "content": exp_prompt},
            ],
            temperature=0.3,
        )
        return resp.choices[0].message.content or ""

    try:
        with ThreadPoolExecutor(max_workers=2) as executor:
            main_future = executor.submit(generate_main_sections)
            exp_future = executor.submit(generate_work_experience)
            raw_main = main_future.result()
            raw_exp = exp_future.result()
    except Exception as e:
        traceback.print_exc()
        return jsonify({"message": f"OpenAI error: {e}"}), 500

    # Merge like new script (append Work Experience at the end), with cleanup
    main_content = clean_markdown(raw_main).strip()
    work_exp_content = clean_markdown(raw_exp).strip()

    if work_exp_content and not work_exp_content.upper().startswith("WORK EXPERIENCE"):
        work_exp_content = "WORK EXPERIENCE\n" + work_exp_content

    merged_text = (main_content + "\n\n" + work_exp_content).strip()
    if not merged_text:
        return jsonify({"message": "Resume generation failed: Empty response"}), 500

    # Save resume content to database
    if candidate_id and job_row_id:
        try:
            job_row = CandidateJob.query.filter_by(id=job_row_id, candidate_id=candidate_id).first()
            if job_row:
                job_row.resume_content = merged_text
                db.session.commit()
        except Exception as e:
            print(f"Warning: Failed to save resume content to database: {e}")

    # File assembly
    candidate_name = merged_text.splitlines()[0].strip()
    safe_name = re.sub(r'[^A-Za-z0-9]+', '_', candidate_name) or "Candidate"

    try:
        if file_type == "word":
            buffer = BytesIO()
            doc = create_resume_word(merged_text)
            doc.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"{safe_name}_resume.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif file_type == "pdf":
            buffer = create_resume_pdf(merged_text)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"{safe_name}_resume.pdf",
                mimetype="application/pdf"
            )
        else:
            return jsonify({"message": "Invalid file_type. Use 'word' or 'pdf'."}), 400
    except Exception as e:
        traceback.print_exc()
        return jsonify({"message": f"File generation error: {e}"}), 500
